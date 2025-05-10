import sys

# noinspection PyUnresolvedReferences
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
# noinspection PyProtectedMember
from docx.table import _Cell
from docx import Document

from docutils import nodes, writers
import logging

logger = logging.getLogger('docx')


def dprint(_func=None, **kw):
    """调试信息打印工具函数
    
    功能：
        - 自动捕获调用者函数名和局部变量
        - 支持关键字参数和自动变量收集两种模式
        - 生成格式化的调试日志信息
    
    参数：
        _func  手动指定的函数名（通常自动获取调用者名称）
        **kw   需要特殊展示的关键字参数对
    
    日志级别：
        - 关键字参数模式使用 INFO 级别
        - 自动变量模式使用 DEBUG 级别（当前实现暂未完全支持）
    """
    # 获取调用栈中上一帧的上下文
    f = sys._getframe(1)
    
    # 处理两种参数模式
    if kw:
        text = ', '.join('%s = %s' % (k, v) for k, v in kw.items())
    else:
        # 自动收集调用者局部变量（排除self）
        text = dict((k, repr(v)) for k, v in f.f_locals.items() if k != 'self')
        text = str(text)

    # 自动获取调用函数名称
    if _func is None:
        _func = f.f_code.co_name

    # 生成日志信息（当前实现统一使用INFO级别）
    if kw:
        logger.info(' '.join([_func, text]))


# noinspection PyUnusedLocal
def _make_depart_admonition(name):
    # noinspection PyMissingOrEmptyDocstring,PyUnusedLocal
    def depart_admonition(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
    return depart_admonition


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring
class DocxWriter(writers.Writer):
    """docutil writer class for docx files"""
    supported = ('docx',)
    settings_spec = ('No options here.', '', ())
    settings_defaults = {}

    output = None
    template_dir = "NO"

    def __init__(self, builder):
        # 初始化父类（docutils的Writer基类）
        writers.Writer.__init__(self)
        self.builder = builder
        # 在调用docx方法前完成模板配置
        self.template_setup()

        # 根据模板配置创建文档对象
        if self.template_dir == "NO":
            dc = Document()  # 创建空白文档
        else:
            dc = Document(self.template_dir)  # 基于模板创建文档
        self.docx_container = dc  # 存储docx文档容器

    def template_setup(self):
        """配置文档模板"""
        # 从Sphinx配置中获取模板路径
        dotx = self.builder.config['docx_template']
        print(dotx)  # 控制台输出模板路径
        if dotx:
            logger.info("MK using template {}".format(dotx))
            self.template_dir = dotx  # 更新模板路径

    def save(self, filename):
        """保存生成的docx文件"""
        self.docx_container.save(filename)  # 调用python-docx的保存方法

    def translate(self):
        """执行文档转换流程"""
        # 创建文档转换器并遍历文档树
        visitor = DocxTranslator(self.document, self.builder, self.docx_container)
        self.document.walkabout(visitor)
        self.output = ''  # 清空输出（实际内容已写入docx_container）


class DocxState:
    """
    文档生成状态跟踪器
    用于管理文档元素（特别是表格中的列表）的嵌套状态
    """
    def __init__(self, location=None):
        self.location = location  # 当前操作位置（文档/表格/单元格）
        self.table = None         # 当前操作的表格对象
        self.column_widths = None # 表格列宽配置
        self.table_style = None   # 表格样式名称
        self.more_cols = 0        # 当前单元格跨越的列数
        self.row = None           # 当前操作的表行
        self.cell_counter = 0     # 当前行中的单元格计数器
        self.ncolumns = 1         # 当前表格的总列数
        "Number of columns in the current table."


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring,PyUnusedLocal
class DocxTranslator(nodes.NodeVisitor):
    """Visitor class to create docx content."""

    def __init__(self, document, builder, docx_container):
        self.builder = builder
        self.docx_container = docx_container
        nodes.NodeVisitor.__init__(self, document)

        # TODO: Perhaps move the list_style into DocxState.
        # However, it should still be a list, and not a separate state,
        # because nested lists are not really nested.
        # So it will only be necessary if there are lists in tables
        # that are in lists.
        self.list_style = []
        self.list_level = 0

        # TODO: And what about sectionlevel?
        self.sectionlevel = 0

        self.table_style_default = 'Grid Table 4'
        self.in_literal_block = False
        self.strong = False
        self.emphasis = False

        self.current_state = DocxState(location=self.docx_container)
        self.current_state.table_style = self.table_style_default

        "The place where paragraphs will be added."
        self.old_states = []
        "A list of older states, e.g. typically [document, table-cell]"

        self.current_paragraph = None
        "The current paragraph that text is being added to."

    def add_text(self, text):
        """向当前段落添加文本内容并应用格式"""
        dprint()
        textrun = self.current_paragraph.add_run(text)  # 创建文本运行对象
        if self.strong:   # 应用粗体格式
            textrun.bold = True
        if self.emphasis:  # 应用斜体格式
            textrun.italic = True

    def new_state(self, location):
        """进入新的文档状态（用于处理嵌套结构）"""
        dprint()
        self.old_states.append(self.current_state)  # 保存当前状态到栈
        self.current_state = DocxState(location=location)  # 创建新状态

    def end_state(self, first=None):
        """退出当前状态，恢复前一个状态"""
        dprint()
        self.current_state = self.old_states.pop()  # 从栈顶弹出历史状态

    def visit_start_of_file(self, node):
        """处理新文件时的初始化（重置章节层级）"""
        dprint()
        # FIXME: 未正确关闭前文件的章节可能导致层级错乱
        self.sectionlevel = 0  # 强制重置标题级别计数器

    def depart_start_of_file(self, node):
        dprint()

    def visit_document(self, node):
        dprint()

    def depart_document(self, node):
        dprint()

    def visit_highlightlang(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_section(self, node):
        dprint()
        self.sectionlevel += 1

    def depart_section(self, node):
        dprint()
        if self.sectionlevel > 0:
            self.sectionlevel -= 1

    def visit_topic(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_topic(self, node):
        dprint()
        raise nodes.SkipNode

    visit_sidebar = visit_topic
    depart_sidebar = depart_topic

    def visit_rubric(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('-[ ')

    def depart_rubric(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(' ]-')

    def visit_compound(self, node):
        dprint()
        pass

    def depart_compound(self, node):
        dprint()
        pass

    def visit_glossary(self, node):
        dprint()
        pass

    def depart_glossary(self, node):
        dprint()
        pass

    def visit_title(self, node):
        dprint()
        self.current_paragraph = self.current_state.location.add_heading(level=self.sectionlevel)

    def depart_title(self, node):
        dprint()

    def visit_subtitle(self, node):
        dprint()
        pass

    def depart_subtitle(self, node):
        dprint()
        pass

    def visit_attribution(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('-- ')

    def depart_attribution(self, node):
        dprint()
        pass

    def visit_desc(self, node):
        dprint()
        pass

    def depart_desc(self, node):
        dprint()
        pass

    def visit_desc_signature(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_desc_signature(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_desc_name(self, node):
        dprint()
        pass

    def depart_desc_name(self, node):
        dprint()
        pass

    def visit_desc_addname(self, node):
        dprint()
        pass

    def depart_desc_addname(self, node):
        dprint()
        pass

    def visit_desc_type(self, node):
        dprint()
        pass

    def depart_desc_type(self, node):
        dprint()
        pass

    def visit_desc_returns(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(' -> ')

    def depart_desc_returns(self, node):
        dprint()
        pass

    def visit_desc_parameterlist(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('(')
        # self.first_param = 1

    def depart_desc_parameterlist(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(')')

    def visit_desc_parameter(self, node):
        dprint()
        raise nodes.SkipNode
        # if not self.first_param:
        #     self.add_text(', ')
        # else:
        #     self.first_param = 0
        # self.add_text(node.astext())
        # raise nodes.SkipNode

    def visit_desc_optional(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('[')

    def depart_desc_optional(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(']')

    def visit_desc_annotation(self, node):
        dprint()
        pass

    def depart_desc_annotation(self, node):
        dprint()
        pass

    def visit_refcount(self, node):
        dprint()
        pass

    def depart_refcount(self, node):
        dprint()
        pass

    def visit_desc_content(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('\n')

    def depart_desc_content(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_figure(self, node):
        # FIXME: figure text become normal paragraph instead of caption.
        dprint()

    def depart_figure(self, node):
        dprint()

    def visit_caption(self, node):
        dprint()
        pass

    def depart_caption(self, node):
        dprint()
        pass

    def visit_productionlist(self, node):
        dprint()
        raise nodes.SkipNode
        # names = []
        # for production in node:
        #     names.append(production['tokenname'])
        # maxlen = max(len(name) for name in names)
        # for production in node:
        #     if production['tokenname']:
        #         self.add_text(production['tokenname'].ljust(maxlen) + ' ::=')
        #         lastname = production['tokenname']
        #     else:
        #         self.add_text('%s    ' % (' '*len(lastname)))
        #     self.add_text(production.astext() + '\n')
        # raise nodes.SkipNode

    def visit_seealso(self, node):
        dprint()

    def depart_seealso(self, node):
        dprint()

    def visit_footnote(self, node):
        dprint()
        raise nodes.SkipNode
        # self._footnote = node.children[0].astext().strip()

    def depart_footnote(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_citation(self, node):
        dprint()
        raise nodes.SkipNode
        # if len(node) and isinstance(node[0], nodes.label):
        #     self._citlabel = node[0].astext()
        # else:
        #     self._citlabel = ''

    def depart_citation(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_label(self, node):
        dprint()
        raise nodes.SkipNode

    # XXX: option list could use some better styling

    def visit_option_list(self, node):
        dprint()
        pass

    def depart_option_list(self, node):
        dprint()
        pass

    def visit_option_list_item(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_option_list_item(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_option_group(self, node):
        dprint()
        raise nodes.SkipNode
        # self._firstoption = True

    def depart_option_group(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('     ')

    def visit_option(self, node):
        dprint()
        raise nodes.SkipNode
        # if self._firstoption:
        #     self._firstoption = False
        # else:
        #     self.add_text(', ')

    def depart_option(self, node):
        dprint()
        pass

    def visit_option_string(self, node):
        dprint()
        pass

    def depart_option_string(self, node):
        dprint()
        pass

    def visit_option_argument(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(node['delimiter'])

    def depart_option_argument(self, node):
        dprint()
        pass

    def visit_description(self, node):
        dprint()
        pass

    def depart_description(self, node):
        dprint()
        pass

    def visit_tabular_col_spec(self, node):
        dprint()
        # TODO: properly implement this!!
        spec = node['spec']
        widths = [float(l.split('cm')[0]) for l in spec.split("{")[1:]]
        self.current_state.column_widths = widths
        raise nodes.SkipNode

    def visit_colspec(self, node):
        dprint()
        # The difficulty here is getting the right column width.
        # This can be specified with a tabular_col_spec, see above.
        #
        # Otherwise it is derived from the number of columns, which is
        # defined in visit_tgroup (a bit hackish).
        # The _block_width is the full width of the document, and this
        # is divided by the number of columns.
        #
        # It would perhaps also be possible to use node['colwidth'] in some way.
        # node['colwidth'] contains an integer like 22, the width of the column in ascii
        if self.current_state.column_widths:
            width = self.current_state.column_widths[0]
            self.current_state.column_widths = self.current_state.column_widths[1:]
            col = self.current_state.table.add_column(Cm(width))
        else:
            # noinspection PyProtectedMember
            col = self.current_state.table.add_column(self.docx_container._block_width // self.current_state.ncolumns)

        raise nodes.SkipNode

    def depart_colspec(self, node):
        dprint()

    def visit_tgroup(self, node):
        dprint()
        colspecs = [c for c in node.children if isinstance(c, nodes.colspec)]
        self.current_state.ncolumns = len(colspecs)

    def depart_tgroup(self, node):
        dprint()
        self.current_state.ncolumns = 1
        pass

    def visit_thead(self, node):
        dprint()
        pass

    def depart_thead(self, node):
        dprint()
        pass

    def visit_tbody(self, node):
        dprint()

    def depart_tbody(self, node):
        dprint()
        pass

    def visit_row(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()
        self.current_state.cell_counter = 0

    def depart_row(self, node):
        dprint()
        pass

    def visit_entry(self, node):
        dprint()
        if 'morerows' in node:
            raise NotImplementedError('Row spanning cells are not implemented.')
        if 'morecols' in node:
            # Hack to make column spanning possible. TODO FIX
            self.current_state.more_cols = node['morecols']

        cell = self.current_state.row.cells[self.current_state.cell_counter]
        # A new paragraph will be added by Sphinx, so remove the automated one
        # This turns out to be not possible, so instead the existing one is
        # reused in visit_paragraph.
        # cell.paragraphs.pop()
        if self.current_state.more_cols:
            # Perhaps this commented line works no too.
            # cell = cell.merge(self.row.cells[self.cell_counter + self.more_cols])
            for i in range(self.current_state.more_cols):
                cell = cell.merge(self.current_state.row.cells[self.current_state.cell_counter + i + 1])

        self.new_state(location=cell)
        # For some annoying reason, a new paragraph is automatically added
        # to each table cell. This is frustrating when you want, e.g. to
        # add a list item instead of a normal paragraph.
        self.current_paragraph = cell.paragraphs[0]

    def depart_entry(self, node):
        dprint()
        self.end_state()
        self.current_state.cell_counter = self.current_state.cell_counter + self.current_state.more_cols + 1
        self.current_state.more_cols = 0

    def visit_table(self, node):
        dprint()

        style = self.current_state.table_style
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.TABLE)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        # Columns are added when a colspec is visited.

        # It is only possible to use a style in add_table when adding a
        # table to the root document. That is, not for a table in a table.
        if len(self.old_states):
            self.current_state.table = self.current_state.location.add_table(rows=0, cols=0)
        else:
            self.current_state.table = self.current_state.location.add_table(
                rows=0, cols=0, style=style)

    def depart_table(self, node):
        dprint()

        self.current_state.table = None
        self.current_state.table_style = self.table_style_default

        # Add an empty paragraph to prevent tables from being concatenated.
        # TODO: Figure out some better solution.
        self.current_state.location.add_paragraph("")

    def visit_acks(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(', '.join(n.astext() for n in node.children[0].children)
        #               + '.')

    def visit_image(self, node):
        dprint()
        uri = node.attributes['uri']
        file_path = f"{self.builder.env.srcdir}/{uri}"
        self.docx_container.add_picture(file_path)  # width=Inches(1.25))
        # .. todo:: 'width' keyword is not supported

    def depart_image(self, node):
        dprint()

    def visit_transition(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('=' * 70)

    def visit_bullet_list(self, node):
        """处理无序列表的进入事件"""
        dprint()
        # TODO: 需要区分编号列表与非编号列表（原设计通过list_style存储样式名）
        # 当前方案通过list_level跟踪嵌套层级（替代原list_style方案）
        # 注意：需与visit_list_item中的样式生成规则保持同步
        self.list_level += 1  # 增加列表嵌套深度计数器

    def depart_bullet_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.append('ListNumber')
        self.list_level += 1

    def depart_enumerated_list(self, node):
        dprint()
        # TODO: self.list_style.pop()
        self.list_level -= 1

    def visit_definition_list(self, node):
        dprint()
        raise nodes.SkipNode
        # self.list_style.append(-2)

    def depart_definition_list(self, node):
        dprint()
        raise nodes.SkipNode
        # self.list_style.pop()

    def visit_list_item(self, node):
        """处理列表项创建的核心逻辑"""
        dprint()
        # 动态生成列表样式名称（根据嵌套层级）
        style = 'List Bullet' if self.list_level < 2 else f'List Bullet {self.list_level}'
        print(f"列表样式 {style}")  # 控制台输出当前列表样式名称
        try:
            # 验证文档是否包含指定列表样式
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            # 样式缺失时的处理（降级为无样式）
            logger.warning(f'列表样式 {style} 缺失，使用默认格式')
            style = None

        # 获取当前操作位置（可能是文档主体或表格单元格）
        curloc = self.current_state.location
        
        # 特殊处理表格单元格中的列表项
        if isinstance(curloc, _Cell):
            # 重用单元格的初始空段落（避免产生多余段落）
            if len(curloc.paragraphs) == 1 and not curloc.paragraphs[0].text:
                self.current_paragraph = curloc.paragraphs[0]
                self.current_paragraph.style = style
            else:
                # 添加新段落并应用列表样式
                self.current_paragraph = curloc.add_paragraph(style=style)
        else:
            # 常规文档流中添加列表段落
            self.current_paragraph = curloc.add_paragraph(style=style)

    def depart_list_item(self, node):
        dprint()

    def visit_definition_list_item(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_definition_list_item(self, node):
        dprint()
        pass

    def visit_term(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_term(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_classifier(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(' : ')

    def depart_classifier(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_definition(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_definition(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_field_list(self, node):
        dprint()
        pass

    def depart_field_list(self, node):
        dprint()
        pass

    def visit_field(self, node):
        dprint()
        pass

    def depart_field(self, node):
        dprint()
        pass

    def visit_field_name(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_field_name(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(':')

    def visit_field_body(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_field_body(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_centered(self, node):
        dprint()
        pass

    def depart_centered(self, node):
        dprint()
        pass

    def visit_hlist(self, node):
        dprint()
        pass

    def depart_hlist(self, node):
        dprint()
        pass

    def visit_hlistcol(self, node):
        dprint()
        pass

    def depart_hlistcol(self, node):
        dprint()
        pass

    def visit_admonition(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_admonition(self, node):
        dprint()
        raise nodes.SkipNode

    def _visit_admonition(self, node):
        dprint()
        raise nodes.SkipNode

    visit_attention = _visit_admonition
    depart_attention = _make_depart_admonition('attention')
    visit_caution = _visit_admonition
    depart_caution = _make_depart_admonition('caution')
    visit_danger = _visit_admonition
    depart_danger = _make_depart_admonition('danger')
    visit_error = _visit_admonition
    depart_error = _make_depart_admonition('error')
    visit_hint = _visit_admonition
    depart_hint = _make_depart_admonition('hint')
    visit_important = _visit_admonition
    depart_important = _make_depart_admonition('important')
    visit_note = _visit_admonition
    depart_note = _make_depart_admonition('note')
    visit_tip = _visit_admonition
    depart_tip = _make_depart_admonition('tip')
    visit_warning = _visit_admonition
    depart_warning = _make_depart_admonition('warning')

    def visit_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
        # if node.children:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + ': ')
        # else:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + '.')

    def depart_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_literal_block(self, node):
        dprint()
        # TODO: Check whether literal blocks work in tables and lists.
        self.in_literal_block = True

        # Unlike with Lists, there will not be a visit to paragraph in a
        # literal block, so we *must* create the paragraph here.
        style = 'Preformatted Text'
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            msg = 'looks like style "{}" is missing\n{}\n using no style'.format(style, repr(exc))
            logger.warning(msg)
            style = None

        self.current_paragraph = self.current_state.location.add_paragraph(style=style)
        self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def depart_literal_block(self, node):
        dprint()
        self.in_literal_block = False

    def visit_doctest_block(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_doctest_block(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_line_block(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_line_block(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_line(self, node):
        dprint()
        pass

    def depart_line(self, node):
        dprint()
        pass

    def visit_block_quote(self, node):
        dprint()

    def depart_block_quote(self, node):
        dprint()

    def visit_compact_paragraph(self, node):
        dprint()

    def depart_compact_paragraph(self, node):
        dprint()

    def visit_paragraph(self, node):
        dprint()

        curloc = self.current_state.location

        if 'List' in self.current_paragraph.style.name and not self.current_paragraph.text:
            # This is the first paragraph in a list item, so do not create another one.
            pass
        elif isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                else:
                    self.current_paragraph = curloc.add_paragraph()
            else:
                self.current_paragraph = curloc.add_paragraph()
            # HACK because the style is messed up, TODO FIX
            self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.current_paragraph.paragraph_format.left_indent = 0
        else:
            self.current_paragraph = curloc.add_paragraph()

    def depart_paragraph(self, node):
        dprint()

    def visit_target(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_index(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_substitution_definition(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_pending_xref(self, node):
        dprint()
        pass

    def depart_pending_xref(self, node):
        dprint()
        pass

    def visit_reference(self, node):
        dprint()
        pass

    def depart_reference(self, node):
        dprint()
        pass

    def visit_download_reference(self, node):
        dprint()
        pass

    def depart_download_reference(self, node):
        dprint()
        pass

    def visit_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = True

    def depart_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = False

    def visit_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def depart_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def visit_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = True

    def depart_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = False

    def visit_abbreviation(self, node):
        dprint()
        # self.add_text('')

    def depart_abbreviation(self, node):
        dprint()
        # if node.hasattr('explanation'):
        #     self.add_text(' (%s)' % node['explanation'])

    def visit_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def depart_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def visit_literal(self, node):
        dprint()
        # self.add_text('``')

    def depart_literal(self, node):
        dprint()
        # self.add_text('``')

    def visit_subscript(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('_')

    def depart_subscript(self, node):
        dprint()
        pass

    def visit_superscript(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('^')

    def depart_superscript(self, node):
        dprint()
        pass

    def visit_footnote_reference(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_citation_reference(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('[%s]' % node.astext())

    def visit_Text(self, node):
        dprint()
        text = node.astext()
        if not self.in_literal_block:
            # assert '\n\n' not in text, 'Found \n\n'
            # Replace double enter with single enter, and single enter with space.
            string_magic = 'TWOENTERSMAGICSTRING'
            text = text.replace('\n\n', string_magic)
            text = text.replace('\n', ' ')
            text = text.replace(string_magic, '\n')
        self.add_text(text)

    def depart_Text(self, node):
        dprint()
        pass

    def visit_generated(self, node):
        dprint()
        pass

    def depart_generated(self, node):
        dprint()
        pass

    def visit_inline(self, node):
        dprint()
        pass

    def depart_inline(self, node):
        dprint()
        pass

    def visit_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('>>')

    def depart_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<<')

    def visit_system_message(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<SYSTEM MESSAGE: %s>' % node.astext())

    def visit_comment(self, node):
        dprint()
        # TODO: FIX Dirty hack / kludge to set table style.
        # Use proper directives or something like that
        comment = node[0]
        if 'DocxTableStyle' in comment:
            self.current_state.table_style = comment.split('DocxTableStyle')[-1].strip()
        raise nodes.SkipNode

    def visit_meta(self, node):
        dprint()
        raise nodes.SkipNode
        # only valid for HTML

    def visit_raw(self, node):
        dprint()
        raise nodes.SkipNode
        # if 'text' in node.get('format', '').split():
        #     self.body.append(node.astext())

    def unknown_visit(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)

    def unknown_departure(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)
